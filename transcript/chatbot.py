# --- START OF MERGED chatbot.py ---

# Imports from both Code 1 and Code 2
from langchain_ollama import ChatOllama
from langchain.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_ollama import OllamaEmbeddings # Included from Code 1, although commented out in Code 2's version
from sentence_transformers import SentenceTransformer
import psycopg2
import re
import argparse
from sqlalchemy import create_engine
from docx import Document # Added for .docx output (from Code 2)
from docx.shared import Pt # Added for font size if needed (optional, from Code 2)
from tqdm import tqdm # Added for progress bar during testing (from Code 2)
import torch

class Chatbot:
    def __init__(self, embedding_model_name="all-MiniLM-L6-v2", llm_model="deepseek-r1:7b", llm_base_url="http://localhost:11434", use_gpu_embeddings=False):
        """
        Initialize the Chatbot with an embedding model and an LLM.
        (Includes initialization prints from Code 2)
        """
        print(f"Initializing Chatbot with embedding model: {embedding_model_name} and LLM: {llm_model}")
        device = "cuda" if use_gpu_embeddings and torch.cuda.is_available() else "cpu"
        self.embedding_model = SentenceTransformer(embedding_model_name, device=device)
        self.embedding_dim = self.embedding_model.get_sentence_embedding_dimension()
        self.llm_model = llm_model
        self.llm_base_url = llm_base_url
        self.llm = ChatOllama(model=self.llm_model, temperature=0.5, base_url=self.llm_base_url)
        print("Chatbot initialized.")

    def get_fragments_from_question(self, question, engine, target_title=None):
        """
        Extract fragments from the PostgreSQL table 'Clips' based on the embedding similarity,
        including their start_time.
        If target_title is provided (from Code 1), only clips with that title are retrieved.
        Uses LIMIT (CONTEXT_LIMIT_AMOUNT) and error handling.
        Returns a list of tuples: (clip_id, transcript, start_time, similarity).
        """
        emb = self.embedding_model.encode(
            question,
            batch_size=32,
            normalize_embeddings=True  # L2 normalization for cosine similarity
        )
        embedding_str = "[" + ",".join(map(str, emb)) + "]"
        conn_str = engine.url.render_as_string(hide_password=False)
        conn = None
        result = []
        CONTEXT_LIMIT_AMOUNT = 4
        try:
            conn = psycopg2.connect(conn_str)
            cur = conn.cursor()

            if target_title:
                # Query now includes c.start_time
                query = f"""
                SELECT c.clip_id, c.transcript, c.start_time,
                    1 - (c.embedding <-> %s) AS similarity
                FROM Clips c
                JOIN originalAudio o ON c.audio_id = o.audio_id
                WHERE o.title = %s
                ORDER BY c.embedding <-> %s
                LIMIT {CONTEXT_LIMIT_AMOUNT}
                """
                cur.execute(query, (embedding_str, target_title, embedding_str))
            else:
                # Query now includes start_time
                query = f"""
                SELECT clip_id, transcript, start_time,
                       1 - (embedding <-> %s) AS similarity
                FROM Clips
                ORDER BY embedding <-> %s
                LIMIT {CONTEXT_LIMIT_AMOUNT}
                """
                cur.execute(query, (embedding_str, embedding_str))

            result = cur.fetchall() # Each row is (clip_id, transcript, start_time, similarity)
            cur.close()
        except psycopg2.Error as e:
            print(f"Database error during fragment retrieval: {e}")
        finally:
            if conn:
                conn.close()
        return result

    def generate_response(self, user_question, engine, conversation_history=[], target_title=None):
        """
        Generate a response by retrieving context from clips (including start_time)
        and feeding it into an LLM.
        The 'documents' string will be formatted to include start_time in MM:SS.
        """

        def format_seconds_to_ms_str(seconds_float):
            """Converts seconds (float) to MM:SS string format."""
            if seconds_float is None:
                return "UNKNOWN" # Should ideally not happen if start_time is NOT NULL
            minutes = int(seconds_float // 60)
            seconds = int(seconds_float % 60) # Truncates fractional seconds
            return f"{minutes:02d}:{seconds:02d}"

        # Retrieve context fragments (now includes start_time)
        # fragments is a list of tuples: (clip_id, transcript, start_time, similarity)
        fragments = self.get_fragments_from_question(user_question, engine, target_title)
        
        formatted_context_blocks = []
        if fragments:
            for i, fragment_data in enumerate(fragments):
                # fragment_data: (clip_id, transcript, start_time, similarity)
                # Indices: transcript=1, start_time=2
                transcript_text = fragment_data[1]
                start_time_seconds = fragment_data[2]

                formatted_time = format_seconds_to_ms_str(start_time_seconds)
                
                # Construct each block as specified:
                # Context X with audio start time <MM:SS>:
                # <Context_X_fragment>
                # ---
                block = f"Context {i+1} with audio start time {formatted_time}:\n{transcript_text}\n---"
                formatted_context_blocks.append(block)
        
        # Join all formatted blocks. If multiple, they are separated by a single newline.
        # This achieves the desired structure:
        # Block1 (ending in ---)
        # Block2 (ending in ---)
        # ...
        documents = "\n".join(formatted_context_blocks)
        context_used_for_response = documents if documents else "No context retrieved."

        # --- END CONTEXT CAPTURE ---

        # print(f"Context retrieved for LLM: {documents[:200]}...") # Optional print from Code 2

        # Build conversation history string up to 2000 characters.
        history_entries = []
        total_length = 0
        for entry in reversed(conversation_history):
            # Ensure keys exist, provide defaults if not (robustness from Code 2)
            user_msg = entry.get('user', '[User message missing]')
            bot_msg = entry.get('chatbot_response', '[Bot response missing]')
            msg = f"User: {user_msg}\nAssistant: {bot_msg}\n{'-'*35}\n"
            if total_length + len(msg) > 2000: # Limit history size
                break
            history_entries.append(msg)
            total_length += len(msg)
        # Reverse back so that messages appear in chronological order
        conversation_history_str = "".join(reversed(history_entries)) if history_entries else "None"

        # Define the prompt including conversation history (using Code 2's template)
        prompt_template = PromptTemplate(
            template="""You are an AI assistant designed to answer questions based *only* on the provided context. The context consists of transcript fragments extracted from an academic lecture audio recording. Each context fragment is presented with its corresponding start time in the audio (e.g., "Context X with audio start time MM:SS: ..."). Utilize both the provided transcript fragments and the conversation history to generate your response.

Important:
* Answer exclusively based on the provided context—do not incorporate outside knowledge or assumptions.
* If the transcript does not contain sufficient information to answer the question, you MUST state, "I do not have enough information in the provided text."
* If the user asks a question in Spanish, you MUST provide your answer in Spanish. Otherwise, respond in English.

Citing Sources:
After you provide your answer, if you used information from the provided context to formulate that answer, you MUST conclude your response by indicating the start time(s) of the specific context fragment(s) you primarily relied upon.
* If you used a single context fragment, state: "(Source: Context at MM:SS)"
* If you used multiple context fragments, state: "(Sources: Context at MM:SS, MM:SS, ...)"
Replace "MM:SS" with the actual start time(s) from the "Context X with audio start time MM:SS:" line of the fragment(s) you used.
* If you state "I do not have enough information in the provided text," do not add this sourcing information.

Conversation History:
{conversation_history}

Provided Context:
{documents}

Question: {user_question}
Answer:
""",
            input_variables=["user_question", "documents", "conversation_history"]
        )
        rag_chain = prompt_template | self.llm | StrOutputParser()

        # print("Invoking RAG chain...") # Optional print from Code 2
        try:
            answer = rag_chain.invoke({
                "user_question": user_question,
                "documents": documents,
                "conversation_history": conversation_history_str,
            })
            # Basic cleaning (remove potential XML-like thinking tags if LLM adds them)
            answer = re.sub(r'<think>.*?</think>', '', answer, flags=re.DOTALL).strip()
            # print("RAG chain invocation complete.") # Optional print from Code 2
        except Exception as e:
            # LLM error handling from Code 2
            print(f"Error invoking LLM chain: {e}")
            answer = "An error occurred while generating the response."

        # Return the standard response structure PLUS the context used (from Code 2)
        return {
            "chatbot_response": answer,
            "context_used": context_used_for_response # Added for testing (from Code 2)
            }

# --- TESTING BLOCK (from Code 2) ---
def run_predefined_tests(chatbot_instance, engine, questions, responses_filename="chatbot_responses.docx", context_filename="context_per_question.docx"):
    """
    Runs a list of predefined questions through the chatbot and saves
    responses and context used into separate .docx files.
    (Function entirely from Code 2)
    """
    print(f"\n--- Starting Predefined Tests ({len(questions)} questions) ---")

    # Initialize .docx documents
    responses_doc = Document()
    context_doc = Document()

    responses_doc.add_heading('Chatbot Test Responses', level=1)
    context_doc.add_heading('Context Used Per Question', level=1)

    for i, question in enumerate(tqdm(questions, desc="Processing Questions")):
        print(f"\nProcessing Question {i+1}/{len(questions)}: {question}")

        # Generate response using empty conversation history
        # NOTE: target_title is not used in test mode based on Code 2's original test setup.
        # If title-specific testing is needed, this call would need modification.
        try:
            response_data = chatbot_instance.generate_response(question, engine, conversation_history=[]) # No target_title here
            chatbot_response = response_data.get("chatbot_response", "Error: Response key not found.")
            context_used = response_data.get("context_used", "Error: Context key not found.")
        except Exception as e:
             print(f"  [ERROR] Failed to get response for question {i+1}: {e}")
             chatbot_response = f"Error during generation: {e}"
             context_used = "N/A due to error"

        # --- Add to Responses Doc ---
        p_resp_q = responses_doc.add_paragraph()
        p_resp_q.add_run(f"{i+1}. ").bold = True
        p_resp_q.add_run(question).bold = True
        responses_doc.add_paragraph(chatbot_response)
        responses_doc.add_paragraph() # Add spacing

        # --- Add to Context Doc ---
        p_cont_q = context_doc.add_paragraph()
        p_cont_q.add_run(f"{i+1}. ").bold = True
        p_cont_q.add_run(question).bold = True
        context_doc.add_paragraph(context_used)
        context_doc.add_paragraph() # Add spacing

    # Save documents
    try:
        responses_doc.save(responses_filename)
        print(f"\nChatbot responses saved to: {responses_filename}")
    except Exception as e:
        print(f"\n[ERROR] Failed to save responses document '{responses_filename}': {e}")

    try:
        context_doc.save(context_filename)
        print(f"Context used saved to: {context_filename}")
    except Exception as e:
        print(f"\n[ERROR] Failed to save context document '{context_filename}': {e}")

    print("--- Predefined Tests Finished ---")


if __name__ == "__main__":
    # Argparse setup from Code 2 (more comprehensive)
    parser = argparse.ArgumentParser(description="Chatbot - Run in interactive mode or execute predefined tests.")
    parser.add_argument("--db_uri", required=True, help="PostgreSQL connection URI (e.g., postgresql://user:pass@host/dbname)")
    parser.add_argument("--mode", choices=['interactive', 'test'], default='interactive', help="Run mode: 'interactive' for live chat, 'test' to run predefined questions.")
    parser.add_argument("--llm", default="deepseek-r1:7b", help="LLM model name (e.g., deepseek-r1:7b, deepseek-r1:14b)") # Allow choosing LLM (from Code 2)
    parser.add_argument("--title", default=None, help="Optional: Filter context by audio title in interactive mode.") # Added to allow using target_title in interactive mode

    args = parser.parse_args()

    print("Creating database engine...")
    try:
        # Engine creation and connection test from Code 2
        engine = create_engine(args.db_uri)
        with engine.connect() as connection:
             print("Database connection successful.")
    except Exception as e:
        print(f"Failed to create database engine or connect: {e}")
        exit(1) # Exit if DB connection fails

    print("Initializing chatbot...")
    # Pass the selected LLM model to the chatbot constructor (from Code 2)
    chatbot = Chatbot(llm_model=args.llm)

    if args.mode == 'test':
        # Test mode execution from Code 2
        # --- List of Predefined Questions (using the second list defined in Code 2) ---
        # --- Ted talk Addiction Questions ---
        # predefined_questions = [
        #     "What personal experience sparked the speaker's interest in addiction?",
        #     "Why did the speaker travel around the world, and who did he meet?",
        #     "What is the traditional view of addiction that the speaker challenges?",
        #     "How does the example of diamorphine (medical heroin) contradict the traditional view of addiction?",
        #     "What was the Rat Park experiment, and what did it show?",
        #     "What human example supports the findings of the Rat Park experiment?",
        #     "What alternative explanation for addiction does the speaker propose?",
        #     "What was Portugal's approach to dealing with addiction, and what were the results?",
        #     "How does our culture typically respond to addiction, according to the speaker?",
        #     "What is the speaker's main message about how we should treat addicts?",
        #     "Why does the speaker criticize the show 'Intervention'?",
        #     "What does the speaker suggest is the true opposite of addiction?",
        #     "How does the speaker relate modern society to the Rat Park experiment?",
        #     "What societal trends does the speaker mention as contributing to disconnection?",
        #     "What did the speaker learn about helping loved ones with addiction?"
        # ]
        # --- Gemini Talk Questions ---
        predefined_questions = [
    "¿Cuál es la novedad que introduce Google con el modelo Gemini Flash 2.0?",
    "¿Cómo se diferencia Gemini Flash 2.0 de otros productos de generación de imágenes existentes?",
    "¿Qué funcionalidades de edición de imágenes ofrece Gemini Flash 2.0?",
    "¿Cuáles son algunos ejemplos prácticos de edición de imágenes presentados en el video?",
    "¿Qué limitaciones se encontraron respecto a la calidad y precisión de las imágenes generadas?",
    "¿Cómo se utiliza Gemini Flash 2.0 a través de Google AI Studio?",
    "¿Qué significa que Gemini Flash 2.0 sea un modelo multimodal?",
    "¿Para qué casos de uso podría ser útil Gemini Flash 2.0 según el video?",
    "¿Cómo se integra el uso de Gemini Flash 2.0 con herramientas de escalado de imágenes?",
    "¿Cuáles son las conclusiones destacadas sobre Gemini Flash 2.0?"
]

        # Run the tests
        run_predefined_tests(chatbot, engine, predefined_questions)

    elif args.mode == 'interactive':
         # Interactive mode execution based on Code 2 (maintains history)
         print("\n--- Starting Interactive Mode (Ctrl+C to exit) ---")
         if args.title:
             print(f"--- Filtering context by title: {args.title} ---")
         conversation_history = [] # Keep track of history in interactive mode
         while True:
            try:
                question = input("\nAsk the chatbot a question: ")
                if question.lower() in ['exit', 'quit']:
                     break
                # Pass current history and get response
                # Pass target_title if provided via command line argument
                response_data = chatbot.generate_response(
                    question,
                    engine,
                    conversation_history=conversation_history,
                    target_title=args.title # Pass title filter if specified
                )
                chatbot_response = response_data.get("chatbot_response", "Sorry, I encountered an issue.")

                print("\nRESPONSE:")
                print(chatbot_response)

                # Add interaction to history for the next turn
                conversation_history.append({
                     "user": question,
                     "chatbot_response": chatbot_response
                })
                # Optional: Prune history if it gets too long over many turns (from Code 2 comment)
                # MAX_HISTORY = 10 # Keep last 10 interactions
                # if len(conversation_history) > MAX_HISTORY:
                #      conversation_history = conversation_history[-MAX_HISTORY:]

            except KeyboardInterrupt:
                print("\nExiting interactive mode.")
                break
            except Exception as e:
                 print(f"\nAn error occurred: {e}")
                 # Optionally reset history or handle error more gracefully

# Example command lines from Code 2 (updated to show --title usage)
# python chatbot.py --db_uri postgresql://admin:secret@localhost:5432/testdb --mode test --llm "gemma3:12b-it-q4_K_M"
# python chatbot.py --db_uri postgresql://admin:secret@localhost:5432/testdb --mode interactive --llm "deepseek-r1:14b"
# python chatbot.py --db_uri postgresql://admin:secret@localhost:5432/testdb --mode interactive --llm "deepseek-r1:7b" --title "My Specific Lecture Title"

# cogito:14b 

# --- END OF MERGED chatbot.py ---